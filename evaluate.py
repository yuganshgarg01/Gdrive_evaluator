"""
evaluate.py
-----------
Evaluates student assignments using Groq AI and fills marks in results.xlsx

Usage:
    python evaluate.py

Folder/file expectations:
    results.xlsx             ← output of download.py
    downloads/               ← all student files (flat folder)
    assignment1.docx         ← assignment questions
    rubric1.docx             ← grading rubric
    assignment2.docx, rubric2.docx  ... and so on

Requirements:
    pip install groq python-docx pymupdf openpyxl pandas
"""

import os
import re
import json
import time
import fitz                          # pymupdf
from docx import Document
from groq import Groq
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

# ============================================================
# CONFIGURATION — Fill these in
# ============================================================

GROQ_API_KEY     = "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"   # https://console.groq.com
GROQ_MODEL       = "llama-3.3-70b-versatile"

RESULTS_EXCEL    = "results.xlsx"
DOWNLOADS_FOLDER = "downloads"
NUM_ASSIGNMENTS  = 3

MAX_STUDENT_TEXT    = 12000   # chars sent to Groq per student
DELAY_BETWEEN_CALLS = 2       # seconds between API calls

# Assignment & rubric files (same folder as this script)
ASSIGNMENT_FILES = {
    1: {"question": "assignment1.docx", "rubric": "rubric1.docx"},
    2: {"question": "assignment2.docx", "rubric": "rubric2.docx"},
    3: {"question": "assignment3.docx", "rubric": "rubric3.docx"},
}

# ============================================================
# STEP 1 — TEXT EXTRACTION (supports PDF, DOCX, TXT, code files)
# ============================================================

CODE_EXTENSIONS = {".py", ".cpp", ".c", ".java", ".js", ".ts", ".html", ".css", ".sql"}


def extract_text_from_pdf(path: str) -> str:
    try:
        doc   = fitz.open(path)
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages).strip()
    except Exception as e:
        return f"[PDF read error: {e}]"


def extract_text_from_docx(path: str) -> str:
    try:
        doc        = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except Exception as e:
        return f"[DOCX read error: {e}]"


def extract_file_text(path: str) -> str:
    """Auto-detect file type and extract text."""
    if not os.path.exists(path):
        return ""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(path)
    elif ext == ".txt":
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception as e:
            return f"[TXT read error: {e}]"
    elif ext in CODE_EXTENSIONS:
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f"[Code: {os.path.basename(path)}]\n```{ext.strip('.')}\n{f.read()}\n```"
        except Exception as e:
            return f"[Code read error: {e}]"
    elif ext in (".jpg", ".jpeg", ".png", ".webp"):
        return f"[Image file: {os.path.basename(path)} — not readable by text model]"
    else:
        # Peek at magic bytes
        with open(path, "rb") as f:
            header = f.read(8)
        if header[:4] == b"%PDF":
            return extract_text_from_pdf(path)
        elif header[:4] == b"PK\x03\x04":
            return extract_text_from_docx(path)
        else:
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read().strip()
            except:
                return f"[Unreadable file: {os.path.basename(path)}]"


def read_student_files(file_paths: list) -> tuple:
    """
    Read one or more files for a student and combine into one text block.
    Returns (combined_text, list_of_filenames)
    """
    all_text  = ""
    filenames = []
    for path in file_paths:
        if not os.path.isfile(path):
            continue
        filename = os.path.basename(path)
        filenames.append(filename)
        text = extract_file_text(path)
        if text:
            all_text += f"\n\n[File: {filename}]\n{text}"
    return all_text.strip(), filenames


# ============================================================
# STEP 2 — PARSE RUBRIC INTO CRITERIA
# ============================================================

def parse_rubric_criteria(rubric_text: str) -> list:
    """
    Parse rubric text into structured criteria list.
    Looks for: "- Criterion Name (10 marks)" or "• Name: 10 pts"
    Falls back to single "Overall" criterion if no pattern found.
    """
    criteria = []
    pattern  = re.findall(
        r"[-•*]?\s*([A-Za-z][^\n(]{2,40})\s*[:(]\s*(\d+)\s*(?:marks?|pts?|points?)",
        rubric_text, re.IGNORECASE
    )
    if pattern:
        for name, pts in pattern:
            criteria.append({"description": name.strip(" :-"), "points": int(pts)})
    else:
        criteria.append({"description": "Overall", "points": 10})
    return criteria


# ============================================================
# STEP 3 — GROQ EVALUATION
# ============================================================

def evaluate_student(client, student_name, filenames, content,
                     question_text, rubric_text, rubric_criteria, assignment_num):

    if not content.strip():
        return build_empty_result(student_name, filenames, "No readable content found", rubric_criteria)

    total_max = sum(c["points"] for c in rubric_criteria)

    # Build criteria section
    criteria_details = ""
    criteria_json    = {}
    for c in rubric_criteria:
        max_pts = c["points"]
        criteria_details += (
            f"\n  Criterion : {c['description']} (max {max_pts} pts)\n"
            f"  Score any INTEGER from 0 to {max_pts}.\n"
            f"  90-100% → Excellent | 75-89% → Good | 50-74% → Average | 30-49% → Below avg | <30% → Poor\n"
        )
        criteria_json[c["description"]] = {
            "score":   f"<integer 0 to {max_pts}>",
            "comment": "<1-2 specific sentences about this student's work>"
        }

    system_prompt = f"""You are a strict university professor evaluating Data Structures assignments.

CRITICAL SCORING RULES:
1. Score CONTINUOUSLY — any integer 0 to max. NO decimals.
2. Every student must get a DIFFERENT score based on actual quality.
3. Base score ONLY on what is written in the submission.
4. Do NOT cluster marks — use the full range.
5. Empty/unreadable submission = 0.

Total max: {total_max} marks.

Respond ONLY in valid JSON — no markdown, no extra text.
Format:
{{
  "criteria": {json.dumps(criteria_json, indent=2)},
  "total": <sum of all scores, integer>,
  "percentage": <total/max*100 to 1 decimal>,
  "grade": "<A>=85%, B>=70%, C>=55%, D>=40%, F<40%>",
  "overall_feedback": "<3-4 lines specific to THIS student>",
  "strengths": "<specific strengths in THIS submission>",
  "improvements": "<what THIS student needs to improve>"
}}"""

    user_message = f"""Assignment {assignment_num} Question:
{question_text}

Rubric:
{rubric_text}

Scoring criteria (use full continuous range):
{criteria_details}

Student: {student_name}
Files: {', '.join(filenames)}

--- SUBMISSION ---
{content[:MAX_STUDENT_TEXT]}{'... [truncated]' if len(content) > MAX_STUDENT_TEXT else ''}
--- END ---

Evaluate {student_name} strictly. Give a UNIQUE score per criterion."""

    try:
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_message}
            ],
            max_tokens=1500,
            temperature=0.1
        )
        raw = response.choices[0].message.content.strip()
        # Strip markdown fences
        if "```" in raw:
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        raw    = raw.strip()
        result = json.loads(raw)
        result["student_name"] = student_name
        result["files"]        = ", ".join(filenames)
        return result

    except json.JSONDecodeError as e:
        print(f"    ✗ JSON parse error ({student_name}): {e}")
        return build_empty_result(student_name, filenames, f"JSON parse error: {e}", rubric_criteria)
    except Exception as e:
        print(f"    ✗ Groq API error ({student_name}): {e}")
        return build_empty_result(student_name, filenames, str(e), rubric_criteria)


def build_empty_result(student_name, filenames, reason, rubric_criteria):
    return {
        "student_name":     student_name,
        "files":            ", ".join(filenames) if filenames else "-",
        "criteria":         {c["description"]: {"score": 0, "comment": reason} for c in rubric_criteria},
        "total":            0,
        "percentage":       0.0,
        "grade":            "F",
        "overall_feedback": reason,
        "strengths":        "-",
        "improvements":     "-"
    }


def format_feedback(result: dict) -> str:
    lines = []
    for cname, cdata in result.get("criteria", {}).items():
        lines.append(f"• {cname}: {cdata.get('score', 0)} pts — {cdata.get('comment', '')}")
    lines.append(f"\nOverall   : {result.get('overall_feedback', '')}")
    lines.append(f"Strengths : {result.get('strengths', '')}")
    lines.append(f"Improve   : {result.get('improvements', '')}")
    lines.append(f"Grade     : {result.get('grade', '-')} ({result.get('percentage', 0)}%)")
    return "\n".join(lines)


# ============================================================
# STEP 4 — FIND STUDENT FILES IN DOWNLOADS FOLDER
# ============================================================

def find_student_files(downloads_folder, rollno, name, assignment_num):
    """
    Find all files for a student + assignment in the downloads folder.
    Supports single file and multiple files per assignment.
    Naming convention: {rollno}_{safe_name}_a{num}.pdf / .docx
    Also matches: {rollno}_{safe_name}_a{num}_1.pdf etc.
    """
    if not os.path.exists(downloads_folder):
        return []
    safe_name = re.sub(r"[^\w\s-]", "", name).strip().replace(" ", "_")
    prefix    = f"{rollno}_{safe_name}_a{assignment_num}".lower()
    matched   = []
    for filename in sorted(os.listdir(downloads_folder)):
        filepath = os.path.join(downloads_folder, filename)
        if not os.path.isfile(filepath):
            continue
        base = os.path.splitext(filename)[0].lower()
        if base.startswith(prefix):
            matched.append(filepath)
    return matched


# ============================================================
# STEP 5 — EXCEL HELPERS
# ============================================================

def find_column(ws, header_name):
    for col in range(1, ws.max_column + 1):
        if ws.cell(row=1, column=col).value == header_name:
            return col
    return None


def style_marks_cell(cell, marks, max_marks):
    try:
        pct = float(marks) / float(max_marks) if max_marks else 0
    except:
        pct = 0
    if pct >= 0.70:
        cell.fill = PatternFill("solid", fgColor="E2EFDA")
    elif pct >= 0.50:
        cell.fill = PatternFill("solid", fgColor="FFF2CC")
    else:
        cell.fill = PatternFill("solid", fgColor="FCE4D6")
    cell.font      = Font(name="Arial", bold=True, size=10)
    cell.alignment = Alignment(horizontal="center", vertical="center")


# ============================================================
# MAIN
# ============================================================

def main():
    print("=" * 55)
    print("  📚 Assignment Evaluator — Powered by Groq")
    print(f"  ⚡ Model : {GROQ_MODEL}")
    print("=" * 55)

    if not GROQ_API_KEY or GROQ_API_KEY == "your_groq_api_key_here":
        print("\nERROR: Set your GROQ_API_KEY at the top of evaluate.py")
        print("Get a free key at: https://console.groq.com")
        return

    if not os.path.exists(RESULTS_EXCEL):
        print(f"\nERROR: '{RESULTS_EXCEL}' not found. Run download.py first.")
        return

    client = Groq(api_key=GROQ_API_KEY)

    # ── Load assignment & rubric texts ─────────────────────────────────────
    print("\nLoading assignment questions and rubrics...")
    assignment_data = {}

    for a_num in range(1, NUM_ASSIGNMENTS + 1):
        q_file = ASSIGNMENT_FILES[a_num]["question"]
        r_file = ASSIGNMENT_FILES[a_num]["rubric"]

        if not os.path.exists(q_file):
            print(f"  WARNING: '{q_file}' not found — skipping Assignment {a_num}")
            continue
        if not os.path.exists(r_file):
            print(f"  WARNING: '{r_file}' not found — skipping Assignment {a_num}")
            continue

        q_text   = extract_text_from_docx(q_file)
        r_text   = extract_text_from_docx(r_file)
        criteria = parse_rubric_criteria(r_text)
        total_max= sum(c["points"] for c in criteria)

        print(f"  ✓ Assignment {a_num}: {len(criteria)} criteria, max {total_max} marks")
        for c in criteria:
            print(f"      - {c['description']}: {c['points']} pts")

        assignment_data[a_num] = {
            "question":  q_text,
            "rubric":    r_text,
            "criteria":  criteria,
            "total_max": total_max
        }

    if not assignment_data:
        print("\nERROR: No assignment/rubric files loaded.")
        return

    # ── Open Excel ─────────────────────────────────────────────────────────
    print(f"\nOpening {RESULTS_EXCEL}...")
    wb = load_workbook(RESULTS_EXCEL)
    ws = wb.active

    col_name  = find_column(ws, "Name")
    col_roll  = find_column(ws, "Roll Number")
    col_total = find_column(ws, "Total Marks")
    col_map   = {}
    for a_num in range(1, NUM_ASSIGNMENTS + 1):
        col_map[f"a{a_num}_path"]     = find_column(ws, f"A{a_num} Local Path")
        col_map[f"a{a_num}_marks"]    = find_column(ws, f"A{a_num} Marks")
        col_map[f"a{a_num}_feedback"] = find_column(ws, f"A{a_num} Feedback")

    total_rows = ws.max_row - 1
    print(f"Evaluating {total_rows} students...\n{'='*55}")

    # ── Evaluate each student ───────────────────────────────────────────────
    for row in range(2, ws.max_row + 1):
        name   = str(ws.cell(row=row, column=col_name).value or f"Row{row}").strip()
        rollno = str(ws.cell(row=row, column=col_roll).value or "").strip()

        print(f"\n[{row-1}/{total_rows}] {name} ({rollno})")

        student_total = 0

        for a_num, data in assignment_data.items():
            marks_col    = col_map.get(f"a{a_num}_marks")
            feedback_col = col_map.get(f"a{a_num}_feedback")
            path_col     = col_map.get(f"a{a_num}_path")

            if not marks_col or not feedback_col:
                continue

            # Skip if already evaluated
            existing = ws.cell(row=row, column=marks_col).value
            if existing not in (None, ""):
                print(f"  ✓ Assignment {a_num}: already done ({existing} marks)")
                try:
                    student_total += float(existing)
                except:
                    pass
                continue

            # Find files: Excel path first, then scan downloads folder
            file_paths = []
            excel_path = ws.cell(row=row, column=path_col).value if path_col else ""
            if excel_path and os.path.exists(str(excel_path)):
                file_paths = [str(excel_path)]
            else:
                file_paths = find_student_files(DOWNLOADS_FOLDER, rollno, name, a_num)

            if not file_paths:
                print(f"  - Assignment {a_num}: no file found → 0 marks")
                ws.cell(row=row, column=marks_col).value    = 0
                ws.cell(row=row, column=feedback_col).value = "File not found or not submitted."
                wb.save(RESULTS_EXCEL)
                continue

            print(f"  - Assignment {a_num}: {len(file_paths)} file(s): {[os.path.basename(p) for p in file_paths]}")

            content, filenames = read_student_files(file_paths)

            result   = evaluate_student(
                client, name, filenames, content,
                data["question"], data["rubric"], data["criteria"], a_num
            )
            marks    = result.get("total", 0)
            max_m    = data["total_max"]
            feedback = format_feedback(result)

            marks_cell = ws.cell(row=row, column=marks_col, value=marks)
            style_marks_cell(marks_cell, marks, max_m)

            fb_cell = ws.cell(row=row, column=feedback_col, value=feedback)
            fb_cell.alignment = Alignment(wrap_text=True, vertical="top")
            fb_cell.font      = Font(name="Arial", size=9)

            student_total += marks
            print(f"    ✓ Marks: {marks}/{max_m} | Grade: {result.get('grade','-')} ({result.get('percentage',0)}%)")

            wb.save(RESULTS_EXCEL)
            time.sleep(DELAY_BETWEEN_CALLS)

        # Write total
        if col_total:
            total_max_all = sum(d["total_max"] for d in assignment_data.values())
            total_cell = ws.cell(row=row, column=col_total, value=student_total)
            style_marks_cell(total_cell, student_total, total_max_all)

        wb.save(RESULTS_EXCEL)

    # ── Final formatting ───────────────────────────────────────────────────
    for a_num in range(1, NUM_ASSIGNMENTS + 1):
        fb_col = col_map.get(f"a{a_num}_feedback")
        if fb_col:
            ws.column_dimensions[get_column_letter(fb_col)].width = 55
    for row in range(2, ws.max_row + 1):
        ws.row_dimensions[row].height = 20

    wb.save(RESULTS_EXCEL)

    # ── Summary ────────────────────────────────────────────────────────────
    totals = []
    if col_total:
        for row in range(2, ws.max_row + 1):
            val = ws.cell(row=row, column=col_total).value
            if val not in (None, ""):
                try:
                    totals.append(float(val))
                except:
                    pass

    print(f"\n{'='*55}")
    print(f"🎉 Done! Results saved to {RESULTS_EXCEL}")
    if totals:
        total_max_all = sum(d["total_max"] for d in assignment_data.values())
        print(f"\n📊 CLASS SUMMARY")
        print(f"   Students  : {len(totals)}")
        print(f"   Average   : {sum(totals)/len(totals):.1f} / {total_max_all}")
        print(f"   Highest   : {max(totals)} / {total_max_all}")
        print(f"   Lowest    : {min(totals)} / {total_max_all}")
    print("=" * 55)


if __name__ == "__main__":
    main()
